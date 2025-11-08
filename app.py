if uploaded_file is not None:
    text = extract_text_from_pdf(uploaded_file)
    st.text_area("识别到的原始文字：", text[:3000])  # 显示前3000字符
import streamlit as st
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from io import BytesIO
from docx import Document
from collections import defaultdict

# 设置中文 OCR 支持
pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

st.title("📄 账单自动整理助手（支持多银行 + OCR）")

st.write("上传你的银行账单（PDF 或 Word），系统会自动识别客户与交易记录并导出 Word 报告。支持扫描账单识别。")

uploaded_file = st.file_uploader("上传账单文件", type=["pdf", "docx"])

def extract_text_from_pdf(pdf_bytes):
    """从 PDF 中提取文本（自动识别文字版或扫描版）"""
    text_content = ""
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text = page.get_text("text")  # 尝试直接提取文本

        # 如果没提取到文字，则使用 OCR
        if not text.strip():
            pix = page.get_pixmap()
            img = Image.open(BytesIO(pix.tobytes("png")))
            ocr_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
            text_content += ocr_text + "\n"
        else:
            text_content += text + "\n"

    return text_content

def group_transactions(text):
    """根据客户名分组交易记录"""
    grouped = defaultdict(list)
    lines = text.splitlines()
    current_name = None

    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 识别客户名（例：客户：张三）
        if any(keyword in line for keyword in ["客户", "名称", "户名"]):
            current_name = line.split(":")[-1].strip()
        elif current_name:
            grouped[current_name].append(line)

    return grouped

def generate_word_report(grouped_data):
    """生成 Word 报告"""
    doc = Document()
    doc.add_heading("转账整理报告", level=1)

    if not grouped_data:
        doc.add_paragraph("未识别到客户或交易记录，请确认账单文字清晰。")
    else:
        for name, records in grouped_data.items():
            doc.add_heading(name, level=2)
            for record in records:
                doc.add_paragraph(record)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if uploaded_file:
    try:
        with st.spinner("正在识别账单内容，请稍候..."):
            text = extract_text_from_pdf(uploaded_file.read())

        if not text.strip():
            st.warning("⚠️ 没有识别到客户或交易记录，请确认账单文字清晰。")
        else:
            grouped_data = group_transactions(text)
            st.success(f"✅ 整理完成，共识别 {len(grouped_data)} 位客户。")

            if st.button("生成 Word 报告"):
                report = generate_word_report(grouped_data)
                st.download_button(
                    label="📥 下载 Word 报告",
                    data=report,
                    file_name="账单整理报告.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

    except Exception as e:
        st.error(f"❌ 发生错误: {e}")
